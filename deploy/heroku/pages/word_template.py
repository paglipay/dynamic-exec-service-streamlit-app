import re
import shutil
import subprocess
import tempfile
import hashlib
from io import BytesIO
from io import StringIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import pandas as pd
import streamlit as st
from docx import Document
from _ai_assistant_panel import render_ai_assistant_panel
from _theme import apply_page_theme


apply_page_theme("Word Template Generator", "Generate filled documents from spreadsheet rows and templates.")
render_ai_assistant_panel("Word Template Generator")
st.write(
    "Upload a spreadsheet and one or more Word templates with placeholders like "
    "<Name> and <Date> to generate filled documents."
)


def load_spreadsheet(uploaded_file) -> pd.DataFrame:
    suffix = Path(uploaded_file.name).suffix.lower()
    if suffix == ".csv":
        return pd.read_csv(uploaded_file)
    if suffix in {".xlsx", ".xls"}:
        return pd.read_excel(uploaded_file)
    raise ValueError("Unsupported spreadsheet format")


def load_pasted_table(pasted_text: str, has_header: bool) -> pd.DataFrame:
    text = pasted_text.strip()
    if not text:
        raise ValueError("No pasted data provided")

    delimiter = "\t" if "\t" in text else ","
    header = 0 if has_header else None
    df = pd.read_csv(StringIO(text), sep=delimiter, header=header)

    if not has_header:
        df.columns = [f"Column{i + 1}" for i in range(df.shape[1])]

    return df


def to_text(value) -> str:
    if pd.isna(value):
        return ""
    return str(value)


def replace_paragraph_text(paragraph, replacements: dict[str, str]) -> None:
    if not paragraph.runs:
        return

    for placeholder, replacement in replacements.items():
        # Keep replacing until this placeholder is no longer present.
        while True:
            full_text = "".join(run.text for run in paragraph.runs)
            start = full_text.find(placeholder)
            if start == -1:
                break

            end = start + len(placeholder)

            spans: list[tuple[int, int, int]] = []
            cursor = 0
            for idx, run in enumerate(paragraph.runs):
                run_end = cursor + len(run.text)
                spans.append((idx, cursor, run_end))
                cursor = run_end

            start_run = None
            end_run = None
            for idx, run_start, run_end in spans:
                if start_run is None and run_start <= start < run_end:
                    start_run = (idx, start - run_start)
                if run_start < end <= run_end:
                    end_run = (idx, end - run_start)
                    break

            # Fallback if a boundary lands exactly on an empty/end boundary.
            if start_run is None or end_run is None:
                break

            start_idx, start_offset = start_run
            end_idx, end_offset = end_run

            if start_idx == end_idx:
                run = paragraph.runs[start_idx]
                run.text = (
                    run.text[:start_offset] + replacement + run.text[end_offset:]
                )
                continue

            first_run = paragraph.runs[start_idx]
            last_run = paragraph.runs[end_idx]

            prefix = first_run.text[:start_offset]
            suffix = last_run.text[end_offset:]

            first_run.text = prefix + replacement
            for idx in range(start_idx + 1, end_idx):
                paragraph.runs[idx].text = ""
            last_run.text = suffix


def apply_replacements(document: Document, replacements: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        replace_paragraph_text(paragraph, replacements)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_paragraph_text(paragraph, replacements)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            replace_paragraph_text(paragraph, replacements)
        for paragraph in section.footer.paragraphs:
            replace_paragraph_text(paragraph, replacements)


def convert_doc_to_docx_bytes(template_bytes: bytes, template_name: str) -> bytes:
    soffice = shutil.which("soffice")
    if not soffice:
        raise RuntimeError(
            "Uploaded .doc templates require LibreOffice (soffice) on the server."
        )

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        src_path = temp_path / template_name
        src_path.write_bytes(template_bytes)

        subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(temp_path),
                str(src_path),
            ],
            check=True,
            capture_output=True,
            text=True,
        )

        converted_path = temp_path / f"{src_path.stem}.docx"
        if not converted_path.exists():
            raise RuntimeError(".doc to .docx conversion failed.")

        return converted_path.read_bytes()


def filled_docx_bytes(template_file, replacements: dict[str, str]) -> bytes:
    suffix = Path(template_file.name).suffix.lower()
    template_bytes = template_file.getvalue()

    if suffix == ".docx":
        input_bytes = template_bytes
    elif suffix == ".doc":
        input_bytes = convert_doc_to_docx_bytes(template_bytes, template_file.name)
    else:
        raise ValueError(f"Unsupported template type: {suffix}")

    document = Document(BytesIO(input_bytes))
    apply_replacements(document, replacements)

    out = BytesIO()
    document.save(out)
    return out.getvalue()


def safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", value.strip())
    cleaned = cleaned.strip("._")
    return cleaned or "output"


def render_filename_pattern(pattern: str, values: dict[str, str]) -> str:
    def replacer(match: re.Match[str]) -> str:
        key = match.group(1)
        return values.get(key, "")

    return re.sub(r"\{([^{}]+)\}", replacer, pattern)


def dataframe_signature(df: pd.DataFrame) -> str:
    sample_csv = df.head(200).to_csv(index=False)
    digest = hashlib.md5(sample_csv.encode("utf-8")).hexdigest()
    return f"{df.shape[0]}x{df.shape[1]}:{digest}"


left_col, right_col = st.columns(2)
with left_col:
    sheet_file = st.file_uploader(
        "Upload spreadsheet (CSV/XLS/XLSX)",
        type=["csv", "xls", "xlsx"],
    )

with right_col:
    pasted_table_text = st.text_area(
        "Or paste spreadsheet cells",
        height=160,
        placeholder=(
            "Select cells in Excel/Sheets, copy, then paste here. "
            "Tab-delimited copied cells are supported."
        ),
    )
    pasted_has_header = st.checkbox(
        "Pasted data includes a header row",
        value=True,
        disabled=not pasted_table_text.strip(),
    )

template_files = st.file_uploader(
    "Upload one or more Word templates (.docx or .doc)",
    type=["docx", "doc"],
    accept_multiple_files=True,
)

show_individual = st.checkbox("Show individual download buttons", value=True)
bundle_zip = st.checkbox("Offer ZIP bundle download", value=True)

if (sheet_file or pasted_table_text.strip()) and template_files:
    try:
        if pasted_table_text.strip():
            df = load_pasted_table(pasted_table_text, pasted_has_header)
        else:
            df = load_spreadsheet(sheet_file)
    except Exception as exc:
        st.error(f"Could not read input data: {exc}")
        st.stop()

    if df.empty:
        st.warning("Spreadsheet has no rows to process.")
        st.stop()

    headers = [str(col) for col in df.columns]
    st.write("### Spreadsheet Preview")
    st.caption("Placeholders are matched exactly as <HeaderName>.")

    if len(headers) != len(set(headers)):
        st.error(
            "Spreadsheet has duplicate column names. "
            "Please make headers unique before generating documents."
        )
        st.stop()

    sig = dataframe_signature(df)
    state_sig_key = "word_template_df_signature"
    state_data_key = "word_template_select_df"

    if st.session_state.get(state_sig_key) != sig:
        st.session_state[state_sig_key] = sig
        st.session_state[state_data_key] = df.copy()
        st.session_state[state_data_key].insert(0, "Include", True)

    select_df: pd.DataFrame = st.session_state[state_data_key]

    action_col1, action_col2 = st.columns(2)
    if action_col1.button("Check all rows"):
        select_df["Include"] = True
    if action_col2.button("Uncheck all rows"):
        select_df["Include"] = False

    st.write("#### Filters")
    filter_columns = st.multiselect(
        "Filter columns",
        options=headers,
        help="Select one or more columns to filter rows before generation.",
    )

    filter_map: dict[str, list[str]] = {}
    for column in filter_columns:
        column_values = sorted(
            {to_text(v) for v in select_df[column].tolist() if to_text(v) != ""}
        )
        selected_values = st.multiselect(
            f"Values for {column}",
            options=column_values,
            default=column_values,
            key=f"filter_values_{column}",
        )
        filter_map[column] = selected_values

    filtered_select_df = select_df
    for column, allowed in filter_map.items():
        if not allowed:
            filtered_select_df = filtered_select_df.iloc[0:0]
            break
        filtered_select_df = filtered_select_df[
            filtered_select_df[column].map(to_text).isin(allowed)
        ]

    edited_filtered_df = st.data_editor(
        filtered_select_df,
        hide_index=True,
        use_container_width=True,
        column_config={
            "Include": st.column_config.CheckboxColumn("Include", default=True),
        },
        disabled=headers,
        key="word_template_filtered_editor",
    )

    for idx in edited_filtered_df.index:
        st.session_state[state_data_key].at[idx, "Include"] = bool(
            edited_filtered_df.at[idx, "Include"]
        )

    selected_count = int(st.session_state[state_data_key]["Include"].sum())
    st.caption(
        f"Selected rows: {selected_count} of {len(st.session_state[state_data_key])}."
    )

    default_name_header = headers[0] if headers else "first_col"
    filename_pattern = st.text_input(
        "Filename pattern",
        value=f"{{template}}_{{{default_name_header}}}_{{row_index}}",
        help=(
            "Use placeholders in braces with spreadsheet headers, for example "
            "{Name}_{Date}. Built-ins: {template}, {row_index}, {first_col}."
        ),
    )

    st.caption(
        "Available header placeholders: "
        + ", ".join(f"{{{header}}}" for header in headers)
    )

    generated_files: list[tuple[str, bytes]] = []
    failed_templates: list[str] = []

    selected_df = st.session_state[state_data_key]
    selected_df = selected_df[selected_df["Include"]].drop(columns=["Include"])
    if filter_columns:
        for column, allowed in filter_map.items():
            if not allowed:
                selected_df = selected_df.iloc[0:0]
                break
            selected_df = selected_df[selected_df[column].map(to_text).isin(allowed)]

    if selected_df.empty:
        st.warning("No rows selected after include checkboxes and filters.")
        st.stop()

    with st.spinner("Generating documents..."):
        for row_index, row in selected_df.iterrows():
            replacements = {f"<{header}>": to_text(row[header]) for header in headers}

            first_col_value = to_text(row.iloc[0])
            base_name = safe_filename(first_col_value) if first_col_value else "row"
            row_tag = f"{base_name}_{row_index + 1}"
            row_values = {header: to_text(row[header]) for header in headers}

            for template_file in template_files:
                template_stem = Path(template_file.name).stem
                name_values = {
                    **row_values,
                    "template": template_stem,
                    "row_index": str(row_index + 1),
                    "first_col": first_col_value,
                }
                rendered_name = render_filename_pattern(filename_pattern, name_values)
                if not rendered_name.strip():
                    rendered_name = f"{template_stem}_{row_tag}"
                out_name = f"{safe_filename(rendered_name)}.docx"

                try:
                    file_bytes = filled_docx_bytes(template_file, replacements)
                    generated_files.append((out_name, file_bytes))
                except Exception as exc:
                    failed_templates.append(
                        f"{template_file.name} (row {row_index + 1}): {exc}"
                    )

    st.success(f"Generated {len(generated_files)} document(s).")

    if failed_templates:
        st.warning("Some templates failed to generate:")
        for item in failed_templates:
            st.write(f"- {item}")

    if show_individual and generated_files:
        st.write("### Individual Downloads")
        for file_name, file_bytes in generated_files:
            st.download_button(
                label=f"Download {file_name}",
                data=file_bytes,
                file_name=file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_{file_name}",
            )

    if bundle_zip and generated_files:
        zip_buffer = BytesIO()
        with ZipFile(zip_buffer, "w", compression=ZIP_DEFLATED) as archive:
            for file_name, file_bytes in generated_files:
                archive.writestr(file_name, file_bytes)

        st.download_button(
            label="Download all as ZIP",
            data=zip_buffer.getvalue(),
            file_name="generated_documents.zip",
            mime="application/zip",
        )
else:
    st.info("Upload a spreadsheet or paste cells, then add at least one template.")
